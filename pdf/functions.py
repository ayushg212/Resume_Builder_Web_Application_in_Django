from docx import Document
from docx.shared import Inches
import json
import os

def create_json(dict):
    json_object = json.dumps(dict, indent=4)
    with open("resume_data.json", "w") as outfile:
        outfile.write(json_object)

def read_json():
    with open('resume_data.json', 'r') as openfile:
        # Reading from json file
        dict = json.load(openfile)
    return dict
 

def create_Doc_1():
    dict = read_json()
    document = Document()
    document.add_heading(dict['name'], level = 0)
    p=document.add_paragraph("Email - ")
    p.add_run(dict['email']).bold = True
    p.add_run("\nPhone - ")
    p.add_run(dict['phone']).bold = True
    document.add_heading("About me", level = 1 )
    document.add_paragraph(dict['summary'])
    document.add_heading("Education Qualificatons", level = 1 )
    edu = ( ('10th',dict['tenth']),
            ( '12th' , dict['twelth'] ),
            ( 'Graduation' , dict['graduation'] ),
            ( 'Post Graduation' , dict['post_graduation'] ) )
    table = document.add_table(rows = 1, cols= 2)
    table.style = 'Colorful List'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Degree"
    hdr_cells[1].text = "Details"
    if dict["post_graduation"] != "" :
        for i in range(4):
            row = table.add_row().cells
            row[0].text = edu[i][0]
            row[1].text = edu[i][1]
    else:
        for i in range(3):
            row = table.add_row().cells
            row[0].text = edu[i][0]
            row[1].text = edu[i][1]
    document.add_heading("Skills", level = 1 )
    document.add_paragraph(dict['skills'])
    document.add_heading("Experience", level = 1 )
    document.add_paragraph(dict['previous_work'])
    document.save('resume1.docx')


def create_Doc_2():
    dict = read_json()
    document = Document()
    a=document.add_heading(dict['name'], level = 0 )
    a.alignment = 2
    p=document.add_paragraph("Email - ")
    p.alignment = 2
    p.add_run(dict['email']).bold = True
    p.add_run("\nPhone - ")
    p.add_run(dict['phone']).bold = True
    document.add_heading("About me", level = 1 )
    document.add_paragraph(dict['summary'])
    document.add_heading("Skills", level = 1 )
    document.add_paragraph(dict['skills'])
    document.add_heading("Experience", level = 1 )
    document.add_paragraph(dict['previous_work'])
    edu = ( ('10th',dict['tenth']),
            ( '12th' , dict['twelth'] ),
            ( 'Graduation' , dict['graduation'] ),
            ( 'Post Graduation' , dict['post_graduation'] ) )
    document.add_heading("Education Qualificatons", level = 1 )
    table = document.add_table(rows = 1, cols= 2)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Degree"
    hdr_cells[1].text = "Details"
    if dict["post_graduation"] != "" :
        for i in range(4):
            row = table.add_row().cells
            row[0].text = edu[i][0]
            row[1].text = edu[i][1]
    else:
        for i in range(3):
            row = table.add_row().cells
            row[0].text = edu[i][0]
            row[1].text = edu[i][1]
    document.save('resume2.docx')


def create_Doc_3():
    dict = read_json()
    document = Document()
    document.add_heading(dict['name'], level = 0 )
    p=document.add_paragraph("Email - ")
    p.add_run(dict['email']).bold = True
    p.add_run("\nPhone - ")
    p.add_run(dict['phone']).bold = True
    document.add_heading("Education Qualificatons", level = 1 )
    document.add_paragraph('10th - ' + dict['tenth'] , style='List Bullet')
    document.add_paragraph('12th - ' + dict['twelth'] , style='List Bullet')
    document.add_paragraph('Graduation - ' + dict['graduation'] , style='List Bullet')
    if dict["post_graduation"] != "" :
        document.add_paragraph('Post Graduation - ' + dict['post_graduation'] , style='List Bullet')
    document.add_heading("About me", level = 1 )
    document.add_paragraph(dict['summary'])
    document.add_heading("Skills", level = 1 )
    document.add_paragraph(dict['skills'])
    document.add_heading("Experience", level = 1 )
    document.add_paragraph(dict['previous_work'])
    document.save('resume3.docx')

def deletefiles():
    if os.path.exists("resume1.docx"):
        os.remove("resume1.docx")
    if os.path.exists("resume1.pdf"):
        os.remove("resume1.pdf")
    if os.path.exists("resume2.docx"):
        os.remove("resume2.docx")
    if os.path.exists("resume2.pdf"):
        os.remove("resume2.pdf")
    if os.path.exists("resume3.docx"):
        os.remove("resume3.docx")
    if os.path.exists("resume3.pdf"):
        os.remove("resume3.pdf")

